import math
from view import db
from docx import Document
from docx.shared import Cm, Mm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from view.models import Question, Team, SubAnswer
from sqlalchemy import func
from datetime import datetime
import os
from docx.enum.style import WD_STYLE_TYPE


answers_per_page = 13
column_width_number = Mm(22.8)
column_width_answer = Mm(162.2)
row_height = Mm(19)
margin = Mm(12.7)
page_width = Mm(210)
page_height = Mm(297)


def create_doc(post):
    # We place all the pubquiz sheets in this folder, with a sub folder for each team
    path = "pubquiz_sheets"
    if not os.path.exists(path):
        os.makedirs(path)
    teams = Team.query.all()
    print("ALLE TEAMS", teams)
    for team in teams:
        path = "pubquiz_sheets" + "/" + team.get_team_name()
        if not os.path.exists(path):
            os.makedirs(path)
        breaks = []
        for p in post:
            breaks.append(p)
        documents = []
        for b in range(0, len(breaks)+1):
            # We create a document for each break (+ 1 because if there are 2 breaks we need 3 separate documents)
            documents.append(Document())
        print("breaks length %s" % len(breaks))
        lines = calculate_lines()

        # We set the sections for each of the documents
        for docu in documents:
            sections = docu.sections
            for section in sections:
                section.top_margin = margin
                section.bottom_margin = margin
                section.left_margin = margin
                section.right_margin = margin
                section.page_width = page_width
                section.page_height = page_height

        # We set the style for each document.
        title_styles = []
        number_cell_styles = []
        for docu in documents:
            styles = docu.styles

            title_style = styles.add_style('TitleCell', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.bold = False
            title_font.size = Pt(22)
            title_style.font.underline = False
            title_styles.append(title_style)

            number_cell_style = styles.add_style('NumberCell', WD_STYLE_TYPE.PARAGRAPH)
            number_cell_font = number_cell_style.font
            number_cell_font.name = 'Calibri'
            number_cell_font.bold = True
            number_cell_font.size = Pt(28)
            number_cell_styles.append(number_cell_style)

        print("VOEG TOE TEAM", team)
        add_team(documents, lines, team, breaks, path, title_styles, number_cell_styles)
        # try:
        #     document.save(path + "/" + team.get_team_name() + "_" + getFileName())
        # except:
        #     return 'Er is iets fout gegaan bij het opslaan van het bestand. Probeer het opnieuw.'

    return 'Documenten zijn opgeslagen'


def calculate_lines():
    amount = 0
    questions = Question.query.all()
    for question in questions:
        if question.questionnumber is not None:
            subanswers = question.subanswers
            for subanswer in subanswers:
                amount = amount + 1
    return amount


def add_team(documents, lines, team, breaks, path, title_styles, number_cell_styles):
    index = 0
    teamname = team.teamname
    pages = math.ceil(lines / answers_per_page)
    fromquestion = 1
    subanswersfromquestion = 0
    roundnumber = 0
    print("AANTAL PAGINAS:", pages)
    while True:
        print("NIEUWE PAGINA VOOR TEAM", teamname, "VANAF VRAAG", fromquestion)
        fromquestion, subanswersfromquestion, roundnumber, lastQ, lastQForBreak = add_page_for_team(documents[index], teamname, fromquestion, subanswersfromquestion, roundnumber, breaks, title_styles[index], number_cell_styles[index])
        if lastQ:
            try:
                # The final question has been determined, we create the final round for the pubquiz
                documents[index].save(path + "/" + team.get_team_name() + "_" + str(index) + "_" + getFileName())
            except:
                return 'Er is iets fout gegaan bij het opslaan van het bestand. Probeer het opnieuw.'
            break
        if lastQForBreak:
            try:
                print("saving sheet for team %s" % team.get_team_name())
                # document.save(path + "/" + team.get_team_name() + "_" + str(index) + "_" + getFileName())
                documents[index].save(path + "/" + team.get_team_name() + "_" + str(index) + "_" + getFileName())
                index += 1
            except:
                return 'Er is iets fout gegaan bij het opslaan van het bestand. Probeer het opnieuw.'
            print("make page")


def add_page_for_team(document, teamname, fromquestion, subanswersfromquestion, roundnumber, breaks, title_style, number_cell_style):
    table = document.add_table(rows=answers_per_page+1, cols=2)
    table.style = 'TableGrid'
    table.rows[0].height = row_height
    title = table.rows[0].cells
    title[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    title[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    title[0].width = column_width_number
    title[1].width = column_width_answer
    title[0].text = 'Quiz'
    title[1].text = 'Naam: ' + teamname
    title[0].paragraphs[0].style = title_style
    title[1].paragraphs[0].style = title_style
    rowsfilled = 0
    title[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    lastQForTeam = False
    lastQForBreak = False
    while rowsfilled < answers_per_page:
        currentrow = table.rows[rowsfilled + 1]
        currentrow.height = row_height
        currentrow.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        currentrow.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        currentrow.cells[0].width = column_width_number
        currentrow.cells[1].width = column_width_answer
        if lastQForTeam:
            rowsfilled = rowsfilled + 1
            continue
        if lastQForBreak:
            rowsfilled = rowsfilled + 1
            continue
        q = Question.query.filter_by(questionnumber=fromquestion).first()
        if q is not None:
            print("test %s and test %s" % (fromquestion, q.id))
            s = db.session.query(func.count(SubAnswer.id)).group_by(SubAnswer.question_id).filter_by(question_id=q.id).all()
            subcount = s[0][0]
            if subcount > subanswersfromquestion:
                if isNewCategory(q, table.rows[rowsfilled]):
                    roundnumber += 1
                    currentrow.cells[0].text = ""
                    currentrow.cells[1].text = "Ronde " + str(roundnumber) + ": " + str(getCategoryNumber(q))
                    currentrow.cells[0].paragraphs[0].style = title_style
                    currentrow.cells[1].paragraphs[0].style = title_style
                    print("NIEUWE RONDE:", roundnumber)
                else:
                    currentrow.cells[0].text = str(q.questionnumber)
                    currentrow.cells[1].text = ""
                    currentrow.cells[0].paragraphs[0].style = number_cell_style
                    currentrow.cells[1].paragraphs[0].style = number_cell_style
                    subanswersfromquestion = subanswersfromquestion + 1
                    print("ZELFDE VRAAG", q.questionnumber)
                currentrow.cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                currentrow.cells[1].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                rowsfilled = rowsfilled + 1
            else:
                print("NIEUWE VRAAG")
                subanswersfromquestion = 0
                fromquestion = fromquestion + 1
        if isLastQuestion(q, subanswersfromquestion):
            print("LAATSTE")
            lastQForTeam = True
        if isBreak(q, subanswersfromquestion, breaks):
            lastQForBreak = True
    return q.questionnumber, subanswersfromquestion, roundnumber, lastQForTeam, lastQForBreak


def getFileName():
    now = datetime.now()
    dt_string = now.strftime("%Y-%m-%d")
    filename = "pubquiz" + dt_string + ".docx"
    index = 1
    while True:
        try:
            open(filename, 'rb')
            filename = "pubquiz" + dt_string + "_" + str(index) + ".docx"
            index += 1
        except:
            return filename
    return


def isNewCategory(question, previousrow):
    if str(question.questionnumber) != str(previousrow.cells[0].text) and (previousrow.cells[0].text != "" or previousrow.cells[0].text == "Quiz"):
        currentCat = question.category_id
        previousQ = Question.query.filter_by(questionnumber=question.questionnumber - 1).first()
        if previousQ:
            previousCat = previousQ.category_id
            if previousCat != currentCat:
                return True
        else:
            return True
    return False


def getCategoryNumber(question):
    return question.questioncategory.name


def isLastQuestion(question, subanswersfromquestion):
    lastQ = Question.query.filter(Question.questionnumber>0).order_by(Question.questionnumber.desc()).first()
    if lastQ.questionnumber == question.questionnumber:
        amountOfSubanswers = db.session.query(func.count(SubAnswer.id)).filter_by(question_id=question.id).all()
        if amountOfSubanswers[0][0] == subanswersfromquestion:
            return True
    return False


def isBreak(question, subanswersfromquestion, breaks):
    if str(question.questionnumber) in breaks:
        print("VRAAGNUMMER", question.questionnumber)
        amountOfSubanswers = db.session.query(func.count(SubAnswer.id)).filter_by(question_id=question.id).all()
        print("BIJ", subanswersfromquestion, "ER ZIJN ER ", amountOfSubanswers[0][0])
        if amountOfSubanswers[0][0] == subanswersfromquestion:
            return True
    return False